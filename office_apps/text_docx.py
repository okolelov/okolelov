from ttkbootstrap import *
from ttkbootstrap.constants import *
import customtkinter as ctk
from threading import Thread
from tkinter import filedialog as fd
import re
import os
from docx import Document

def rgb(rgb):
	return "#%02x%02x%02x" % rgb

normal = rgb((234, 234, 234))
keywords =rgb((234, 95, 95))
comments = rgb((95, 234, 165))
background = rgb((42, 42, 42))
tfont = ('Times New Roman', '15', 'bold')

class App(Window):
	def __init__(self):
		super().__init__(title="Ruffice документ", themename="darkly")
		Thread(target=self.user_interface).start()

	def user_interface(self):
		self.w, self.h = self.winfo_screenwidth(), self.winfo_screenheight()
		self.geometry("%dx%d" % (self.w, self.h))
		self.attributes('-zoomed', True)
		self.frame = Frame(
			self,
			width=1000
			)
		self.frame.pack(expand=YES, fill="y")
		self.frame.pack_propagate(False)
		self.text_fild = ctk.CTkTextbox(
			self.frame,
			bg_color = normal,
			wrap="word",
			spacing3 = 10,
			font=('Times New Roman', 16, 'normal')
			)
		self.text_fild.pack(expand=YES,fill="both")

		self.file = Menu(self)  
		self.new_item = Menu(self.file, tearoff=False)  
		self.new_item.add_command(label='Сохранить', command=self.save_file)
		self.new_item.add_command(label='Загрузить', command=self.load_file)
		self.new_item.add_command(label='Выход', command=self.destroy)  
		self.file.add_cascade(label='Файл', menu=self.new_item) 
		self.configure(menu=self.file)

	def load_file(self):
		self.file_name = fd.askopenfilename()
		doc = Document(self.file_name)
		app.file_load=self.file_name
		text=[]
		self.text_fild.delete("0.0", "end")
		count = len(doc.paragraphs)
		for i in range(count):
			n = doc.paragraphs[i].text
			text.append(n)
		for i in range(len(text)):
			a = text[i]
			self.text_fild.insert(END, a)

	def save_file(self):
		doc = Document(self.file_name)
		count = len(doc.paragraphs)
		for i in range(count):
			n = doc.paragraphs[i].text
			ar = n+"\n"	
			doc.paragraphs[i].clear()
		text=self.text_fild.get("0.0","end")
		doc.add_paragraph(text)
		doc.save(self.fileload)

app = App()
app.mainloop()
		